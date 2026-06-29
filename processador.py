"""
processador.py
- Lê planilha Excel do DP e retorna lista de funcionários
- Substitui variáveis nos modelos .docx
- Converte .docx → PDF via LibreOffice
"""

import os
import re
import shutil
import subprocess
from datetime import date
from pathlib import Path

import openpyxl
from docx import Document

from config import MODELOS_DIR, OUTPUT_DIR, EMPRESA, RESP_SST, CNPJ, DOCUMENTOS


# ══════════════════════════════════════════════════════════
#  LEITURA DA PLANILHA
# ══════════════════════════════════════════════════════════

# Mapeamento: nome da coluna no Excel → chave interna
COLUNAS_MAP = {
    "nome":        ["nome", "funcionário", "funcionario", "colaborador"],
    "cpf":         ["cpf"],
    "matricula":   ["matrícula", "matricula", "mat", "esocial matricula", "esocial matrícula", "esocial"],
    "cargo":       ["cargo atual", "cargo", "função", "funcao", "função atual", "código atual", "codigo atual"],
    "lotacao":     ["contrato", "lotação", "lotacao", "obra", "frente"],
    "admissao":    ["admissão", "admissao", "data admissão", "data de admissão", "dt admissão", "admissão"],
    "celular":     ["celular", "telefone", "whatsapp", "fone", "cel"],
    "email":       ["e-mail pessoal", "email", "e-mail", "email pessoal"],
    "cbo":         ["cbo", "código cbo", "codigo cbo", "cbo cargo", "cbo função", "cbo funcao"],
}


def _mapear_cabecalho(headers: list[str]) -> dict:
    """Mapeia posição das colunas do Excel para chaves internas."""
    mapa = {}
    for i, h in enumerate(headers):
        if h is None:
            continue
        h_norm = str(h).strip().lower()
        for chave, aliases in COLUNAS_MAP.items():
            if h_norm in aliases and chave not in mapa:
                mapa[chave] = i
    return mapa


def ler_planilha(caminho: str) -> tuple[list[dict], list[str]]:
    """
    Lê planilha Excel do DP.
    Retorna (lista_funcionarios, avisos).
    """
    wb = openpyxl.load_workbook(caminho, read_only=True, data_only=True)
    try:
        ws = wb.active

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return [], ["Planilha vazia."]

        # Encontra linha de cabeçalho (primeira linha não vazia)
        header_row = None
        header_idx = 0
        for i, row in enumerate(rows):
            if any(cell is not None for cell in row):
                header_row = [str(c).strip() if c is not None else "" for c in row]
                header_idx = i
                break

        if header_row is None:
            return [], ["Não foi possível identificar o cabeçalho."]

        mapa = _mapear_cabecalho(header_row)
        avisos = []

        if "nome" not in mapa:
            avisos.append("⚠️ Coluna 'Nome' não encontrada.")
        if "cpf" not in mapa:
            avisos.append("⚠️ Coluna 'CPF' não encontrada.")

        funcionarios = []
        for row in rows[header_idx + 1:]:
            if not any(cell is not None for cell in row):
                continue

            def get(chave):
                idx = mapa.get(chave)
                if idx is None or idx >= len(row):
                    return ""
                v = row[idx]
                return str(v).strip() if v is not None else ""

            nome = get("nome")
            cpf  = get("cpf")

            if not nome or not cpf:
                continue

            # Formata CPF
            cpf_num = re.sub(r"\D", "", cpf)
            if len(cpf_num) == 11:
                cpf = f"{cpf_num[:3]}.{cpf_num[3:6]}.{cpf_num[6:9]}-{cpf_num[9:]}"

            # Formata celular
            cel = re.sub(r"\D", "", get("celular"))

            funcionarios.append({
                "nome":      nome,
                "cpf":       cpf,
                "matricula": get("matricula"),
                "cargo":     get("cargo"),
                "lotacao":   get("lotacao"),
                "admissao":  get("admissao"),
                "celular":   cel,
                "email":     get("email"),
                "cbo":       re.sub(r"\D", "", get("cbo")),  # só dígitos
            })

        return funcionarios, avisos
    finally:
        wb.close()


# ══════════════════════════════════════════════════════════
#  PREENCHIMENTO DO DOCX
# ══════════════════════════════════════════════════════════

def _substituir_texto(texto: str, variaveis: dict) -> str:
    """Substitui {{VARIAVEL}} e @variavel pelo valor correspondente."""
    for chave, valor in variaveis.items():
        texto = texto.replace(f"{{{{{chave}}}}}", str(valor) if valor else "")
    # Também substitui formato @variavel (usado pelos documentos de treinamento)
    for chave, valor in variaveis.items():
        texto = texto.replace(f"@{chave.lower()}", str(valor) if valor else "")
    return texto


def _processar_paragrafo(paragrafo, variaveis: dict):
    """Preserva formatação ao substituir variáveis no parágrafo."""
    # Primeiro tenta substituição simples em cada run
    for run in paragrafo.runs:
        if "{{" in run.text or "@" in run.text:
            run.text = _substituir_texto(run.text, variaveis)

    # Se ainda há variável espalhada entre runs, une e redistribui
    texto_completo = "".join(r.text for r in paragrafo.runs)
    if "{{" in texto_completo or "@" in texto_completo:
        texto_novo = _substituir_texto(texto_completo, variaveis)
        if paragrafo.runs:
            paragrafo.runs[0].text = texto_novo
            for run in paragrafo.runs[1:]:
                run.text = ""


def preencher_docx(modelo_id: str, funcionario: dict, pasta_saida: str) -> str | None:
    """
    Preenche um modelo .docx com os dados do funcionário.
    Retorna o caminho do arquivo gerado ou None se modelo não existe.
    """
    os.makedirs(pasta_saida, exist_ok=True)

    # Monta dicionário de variáveis
    variaveis = {
        "NOME":          funcionario.get("nome", ""),
        "CPF":           funcionario.get("cpf", ""),
        "MATRICULA":     funcionario.get("matricula", funcionario.get("cpf", "")),
        "CARGO":         funcionario.get("cargo", ""),
        "LOTACAO":       funcionario.get("lotacao", ""),
        "DATA_ADMISSAO": funcionario.get("admissao", ""),
        "DATA_HOJE":     date.today().strftime("%d/%m/%Y"),
        "CELULAR":       funcionario.get("celular", ""),
        "EMAIL":         funcionario.get("email", ""),
        "EMPRESA":       EMPRESA,
        "RESP_SST":      RESP_SST,
        "CNPJ":          CNPJ,
        # Aliases para formato @variavel usado nos treinamentos
        "funcao":        funcionario.get("cargo", ""),
        "dt_adm":        funcionario.get("admissao", ""),
        "resp_tecnico":  RESP_SST,
        "lotacao":       funcionario.get("lotacao", ""),
        "matricula":     funcionario.get("matricula", funcionario.get("cpf", "")),
        "ctps":          funcionario.get("ctps", ""),
        "rg":            funcionario.get("rg", ""),
        "cpf":           funcionario.get("cpf", ""),
        "nome":          funcionario.get("nome", ""),
        "empresa":       EMPRESA,
        "cargo":         funcionario.get("cargo", ""),
        "data_hoje":     date.today().strftime("%d/%m/%Y"),
    }

    # Tenta banco primeiro, fallback para disco
    import banco as _banco
    import io
    cargo_func = funcionario.get("cargo", "")

    # Ficha de EPI: SEMPRE gera dinamicamente com os EPIs atuais do cargo
    # (nunca usa ficha pré-salva que pode estar desatualizada)
    if modelo_id == "10_ficha_controle_epi":
        # Busca template base (cargo=None) para usar como molde
        base_bytes = _banco.buscar_modelo("10_ficha_controle_epi")
        if not base_bytes:
            _disco_epi = os.path.join(MODELOS_DIR, "10_ficha_controle_epi.docx")
            if os.path.exists(_disco_epi):
                with open(_disco_epi, "rb") as _f:
                    base_bytes = _f.read()
        # Fallback: usa qualquer ficha cargo-específica como template base
        if not base_bytes:
            _todos = _banco.listar_modelos()
            for _m in _todos:
                if _m["id"] == "10_ficha_controle_epi" and _m.get("cargo") and _m.get("tem_conteudo"):
                    base_bytes = _banco.buscar_modelo("10_ficha_controle_epi", cargo=_m["cargo"])
                    if base_bytes:
                        print(f"  ⚠️  [{modelo_id}] usando ficha de cargo={_m['cargo']!r} como template base")
                        break
        if not base_bytes:
            print(f"  ❌ [{modelo_id}] template base não encontrado em nenhuma fonte")
            return None
        try:
            epis_list = _banco.listar_epis_do_cargo(cargo_func) or []
            if epis_list:
                conteudo_banco = preencher_ficha_epi_dinamica(funcionario, epis_list, base_bytes)
                print(f"  ✅ [{modelo_id}] Ficha EPI gerada com {len(epis_list)} EPIs do cargo={cargo_func!r}")
            else:
                conteudo_banco = base_bytes
                print(f"  ⚠️  [{modelo_id}] sem EPIs cadastrados para cargo={cargo_func!r}, usando template base")
        except Exception as _e:
            conteudo_banco = base_bytes
            print(f"  ⚠️  [{modelo_id}] erro ao gerar EPI dinâmica: {_e}")
        import io as _io
        doc = Document(_io.BytesIO(conteudo_banco))
        for para in doc.paragraphs:
            _processar_paragrafo(para, variaveis)
        for section in doc.sections:
            for para in section.header.paragraphs:
                _processar_paragrafo(para, variaveis)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _processar_paragrafo(para, variaveis)
            for para in section.footer.paragraphs:
                _processar_paragrafo(para, variaveis)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _processar_paragrafo(para, variaveis)
        buf = _io.BytesIO()
        doc.save(buf)
        caminho = os.path.join(pasta_saida, f"{modelo_id}.docx")
        with open(caminho, "wb") as _fout:
            _fout.write(buf.getvalue())
        return caminho

    conteudo_banco = _banco.buscar_modelo(modelo_id, cargo=cargo_func)

    # Fallback especial: gera OS a partir do 03_os_base quando não há OS salvo para o cargo
    if not conteudo_banco and modelo_id == "03_os":
        base_bytes = _banco.buscar_modelo("03_os_base")
        if base_bytes:
            try:
                pgr = _banco.buscar_pgr_cargo(cargo_func) or {}
                epis_list = _banco.listar_epis_do_cargo(cargo_func) or []
                epis_texto = "\n".join(
                    f"- {e.get('descricao','')} (CA: {e.get('ca','')})"
                    for e in epis_list
                ) or "—"
                conteudo_banco = preencher_os_dinamica(
                    funcionario=funcionario,
                    descricao_atividades=pgr.get("atividades", ""),
                    epis_texto=epis_texto,
                    modelo_bytes=base_bytes,
                    riscos_texto=pgr.get("riscos", ""),
                )
                print(f"  ✅ [{modelo_id}] OS gerada automaticamente do base para cargo={cargo_func!r}")
            except Exception as _e:
                print(f"  ⚠️  [{modelo_id}] falha ao gerar OS automática: {_e}")

    if conteudo_banco:
        print(f"  ✅ [{modelo_id}] usando banco ({len(conteudo_banco)} bytes, cargo={cargo_func!r})")
        doc = Document(io.BytesIO(conteudo_banco))
    else:
        modelo_path = os.path.join(MODELOS_DIR, f"{modelo_id}.docx")
        if not os.path.exists(modelo_path):
            print(f"  ❌ [{modelo_id}] não encontrado no banco nem no disco")
            return None
        print(f"  ⚠️  [{modelo_id}] usando disco (não está no banco)")
        doc = Document(modelo_path)

    # Processa parágrafos do corpo
    for para in doc.paragraphs:
        _processar_paragrafo(para, variaveis)

    # Processa tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for para in celula.paragraphs:
                    _processar_paragrafo(para, variaveis)

    # Processa cabeçalho e rodapé
    for section in doc.sections:
        for para in section.header.paragraphs:
            _processar_paragrafo(para, variaveis)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _processar_paragrafo(para, variaveis)
        for para in section.footer.paragraphs:
            _processar_paragrafo(para, variaveis)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _processar_paragrafo(para, variaveis)

    # Nome seguro para o arquivo
    nome_seguro = re.sub(r"[^\w\s-]", "", funcionario.get("nome", "funcionario"))
    nome_seguro = re.sub(r"\s+", "_", nome_seguro.strip())
    nome_arquivo = f"{modelo_id}__{nome_seguro}.docx"
    caminho_docx = os.path.join(pasta_saida, nome_arquivo)

    doc.save(caminho_docx)
    return caminho_docx


# ══════════════════════════════════════════════════════════
#  CONVERSÃO DOCX → PDF
# ══════════════════════════════════════════════════════════

def _libreoffice_path() -> str | None:
    """Localiza o executável do LibreOffice no Windows ou Linux."""
    candidatos = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
        "libreoffice",
    ]
    for c in candidatos:
        if os.path.exists(c):
            return c
        if shutil.which(c):
            return c
    return None


def converter_para_pdf(caminho_docx: str) -> str | None:
    """
    Converte .docx para PDF usando LibreOffice headless.
    Retorna caminho do PDF gerado ou None em caso de erro.
    """
    soffice = _libreoffice_path()
    if not soffice:
        print("❌ LibreOffice não encontrado. Instale em https://www.libreoffice.org")
        return None

    pasta = os.path.dirname(caminho_docx)
    try:
        resultado = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", pasta, caminho_docx],
            capture_output=True, text=True, timeout=60
        )
        if resultado.returncode != 0:
            print(f"❌ Erro ao converter: {resultado.stderr}")
            return None

        caminho_pdf = caminho_docx.replace(".docx", ".pdf")
        if os.path.exists(caminho_pdf):
            return caminho_pdf

        # LibreOffice às vezes gera nome diferente
        base = Path(caminho_docx).stem
        for f in os.listdir(pasta):
            if f.endswith(".pdf") and base in f:
                return os.path.join(pasta, f)

        return None
    except subprocess.TimeoutExpired:
        print("❌ Timeout ao converter para PDF.")
        return None
    except Exception as e:
        print(f"❌ Erro inesperado: {e}")
        return None


# ══════════════════════════════════════════════════════════
#  FLUXO COMPLETO: gerar kit de um funcionário
# ══════════════════════════════════════════════════════════

def gerar_kit_funcionario(funcionario: dict, doc_ids: list[str], lote_pasta: str) -> list[dict]:
    """
    Gera todos os PDFs do kit de um funcionário.
    Retorna lista de dicts com {doc_id, doc_nome, pdf_path, erro}.
    """
    nome_seguro = re.sub(r"[^\w\s-]", "", funcionario.get("nome", "func"))
    nome_seguro = re.sub(r"\s+", "_", nome_seguro.strip())
    pasta_func  = os.path.join(lote_pasta, nome_seguro)
    os.makedirs(pasta_func, exist_ok=True)

    resultados = []
    docs_map   = {d["id"]: d["nome"] for d in DOCUMENTOS}

    for doc_id in doc_ids:
        doc_nome = docs_map.get(doc_id, doc_id)
        print(f"  📄 {doc_nome} → {funcionario['nome']}")

        caminho_docx = preencher_docx(doc_id, funcionario, pasta_func)
        if not caminho_docx:
            resultados.append({"doc_id": doc_id, "doc_nome": doc_nome,
                                "pdf_path": None, "erro": "Modelo não encontrado"})
            continue

        caminho_pdf = converter_para_pdf(caminho_docx)
        if not caminho_pdf:
            resultados.append({"doc_id": doc_id, "doc_nome": doc_nome,
                                "pdf_path": None, "erro": "Falha na conversão PDF"})
            continue

        # Remove o .docx intermediário (mantém só o PDF)
        try:
            os.remove(caminho_docx)
        except Exception:
            pass

        resultados.append({"doc_id": doc_id, "doc_nome": doc_nome,
                            "pdf_path": caminho_pdf, "erro": None})

    return resultados


def preencher_ficha_epi_dinamica(funcionario: dict, epis: list, modelo_bytes: bytes) -> bytes:
    """
    Preenche a ficha de EPI base com os EPIs selecionados dinamicamente.
    epis = lista de dicts com {descricao, ca, quantidade}
    Retorna bytes do .docx preenchido.
    """
    import io
    from datetime import date

    doc = Document(io.BytesIO(modelo_bytes))

    variaveis = {
        "NOME":          funcionario.get("nome", ""),
        "CPF":           funcionario.get("cpf", ""),
        "MATRICULA":     funcionario.get("matricula", funcionario.get("cpf", "")),
        "CARGO":         funcionario.get("cargo", ""),
        "LOTACAO":       funcionario.get("lotacao", ""),
        "DATA_ADMISSAO": funcionario.get("admissao", ""),
        "DATA_HOJE":     date.today().strftime("%d/%m/%Y"),
        "CELULAR":       funcionario.get("celular", ""),
        "EMAIL":         funcionario.get("email", ""),
        "EMPRESA":       EMPRESA,
        "RESP_SST":      RESP_SST,
        "CNPJ":          CNPJ,
        "CTPS":          funcionario.get("ctps", ""),
        "RG":            funcionario.get("rg", ""),
        "TITULO_FICHA":  "FICHA DE ENTREGA DE EPI/EPC/UNIFORMES",
        # Aliases para templates antigos com @funcao, @dt_adm, @nome, etc.
        "FUNCAO":        funcionario.get("cargo", ""),
        "DT_ADM":        funcionario.get("admissao", ""),
        "ADMISSAO":      funcionario.get("admissao", ""),
        "FUNCIONARIO":   funcionario.get("nome", ""),
        "CPF_FUNC":      funcionario.get("cpf", ""),
        "MAT":           funcionario.get("matricula", funcionario.get("cpf", "")),
    }

    # Preenche variáveis nos parágrafos e cabeçalho/rodapé
    for para in doc.paragraphs:
        _processar_paragrafo(para, variaveis)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _processar_paragrafo(para, variaveis)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _processar_paragrafo(para, variaveis)
        for para in section.footer.paragraphs:
            _processar_paragrafo(para, variaveis)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _processar_paragrafo(para, variaveis)

    # Encontra a tabela principal e preenche as linhas de EPI
    import copy
    data_hoje = date.today().strftime("%d/%m/%Y")
    WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _celulas_unicas(row):
        seen, cells = set(), []
        for cell in row.cells:
            if id(cell._tc) not in seen:
                seen.add(id(cell._tc))
                cells.append(cell)
        return cells

    def _escrever_celula(cell, texto):
        """Zera todos os runs da célula e escreve apenas o texto desejado."""
        if not cell.paragraphs:
            return
        para = cell.paragraphs[0]
        # Esvazia todos os runs existentes
        for run in para.runs:
            run.text = ""
        # Escreve no primeiro run, ou cria um novo
        if para.runs:
            para.runs[0].text = str(texto) if texto else ""
        elif texto:
            para.add_run(str(texto))

    def _clonar_linha(table, row_ref):
        new_tr = copy.deepcopy(row_ref._tr)
        # Limpa todos os <w:t> na cópia
        for t_el in new_tr.findall(f".//{WNS}t"):
            t_el.text = ""
        table._tbl.append(new_tr)
        return table.rows[-1]

    for table in doc.tables:
        # Identifica a linha de cabeçalho (ENTREGA / DESCRIÇÃO / C.A / QUANTIDADE)
        header_row = None
        for ri, row in enumerate(table.rows):
            celulas_txt = [c.text.strip().upper() for c in row.cells]
            if any("ENTREGA" in c or "DESCRI" in c or "C.A" in c for c in celulas_txt):
                header_row = ri
                break

        # Fallback: tabela com 4+ colunas
        if header_row is None:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                header_row = 0

        if header_row is None:
            continue

        # Mapeia colunas pelo texto do cabeçalho
        header_cells = _celulas_unicas(table.rows[header_row])
        col_entrega = col_desc = col_ca = col_qtd = None
        for ci, cell in enumerate(header_cells):
            t = cell.text.strip().upper()
            if "ENTREGA" in t or ("DATA" in t and col_entrega is None):
                col_entrega = ci
            elif "DESCRI" in t:
                col_desc = ci
            elif "C.A" in t or t == "CA":
                col_ca = ci
            elif "QUANT" in t:
                col_qtd = ci

        # Fallback posicional se não detectou
        if col_entrega is None: col_entrega = 0
        if col_desc    is None: col_desc    = 1
        if col_ca      is None: col_ca      = 2
        if col_qtd     is None: col_qtd     = 3

        print(f"[ficha_epi] colunas: entrega={col_entrega} desc={col_desc} ca={col_ca} qtd={col_qtd}")

        # Localiza linhas de dados (entre cabeçalho e seção de assinatura)
        sig_keywords = ("assinatura", "declaro", "ciente", "___", "funcionário:", "trabalhador")
        epi_rows = []
        for ri in range(header_row + 1, len(table.rows)):
            texto = " ".join(c.text.strip().lower() for c in table.rows[ri].cells)
            if any(kw in texto for kw in sig_keywords):
                break
            epi_rows.append(ri)

        # Clona linhas extras se necessário
        row_template = table.rows[epi_rows[-1]] if epi_rows else table.rows[header_row]
        while len(epi_rows) < len(epis):
            _clonar_linha(table, row_template)
            epi_rows.append(len(table.rows) - 1)

        epi_rows_set = set(epi_rows)

        # Preenche cada linha de EPI
        for i, ri in enumerate(epi_rows):
            row = table.rows[ri]
            cells = _celulas_unicas(row)
            if i < len(epis):
                epi = epis[i]
                vals = {
                    col_entrega: data_hoje,
                    col_desc:    epi.get("descricao", ""),
                    col_ca:      str(epi.get("ca", "") or ""),
                    col_qtd:     str(epi.get("quantidade", 1)),
                }
                for ci, cell in enumerate(cells):
                    _escrever_celula(cell, vals.get(ci, ""))
            else:
                # Linha extra vazia — limpa
                for cell in cells:
                    _escrever_celula(cell, "")

        # Preenche variáveis nas demais linhas (NOME, EMPRESA etc.)
        for ri, row in enumerate(table.rows):
            if ri in epi_rows_set:
                continue
            seen_cell = set()
            for cell in row.cells:
                if id(cell._tc) in seen_cell:
                    continue
                seen_cell.add(id(cell._tc))
                for para in cell.paragraphs:
                    _processar_paragrafo(para, variaveis)

        print(f"[ficha_epi] tabela preenchida: {len(epis)} EPI(s) em {len(epi_rows)} linha(s)")
        break  # Processa só a primeira tabela principal
    else:
        if epis:
            print(f"[ficha_epi] AVISO: nenhuma tabela compatível no template — {len(epis)} EPI(s) não preenchidos")
            for ti, t in enumerate(doc.tables):
                print(f"  tabela[{ti}]: {len(t.rows)}x{len(t.columns)} | header={[c.text.strip() for c in t.rows[0].cells]}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def preencher_os_dinamica(funcionario: dict, descricao_atividades: str,
                          epis_texto: str, modelo_bytes: bytes,
                          riscos_texto: str = "") -> bytes:
    """
    Preenche a OS modelo com dados do funcionário + CBO + EPIs.
    - descricao_atividades: texto da descrição sumária do CBO
    - epis_texto: lista de EPIs formatada (ex: "- BOTINA DE SEGURANÇA (CA: 48413)\n- ...")
    - modelo_bytes: bytes do modelo 03_os_base.docx
    Retorna bytes do .docx preenchido.
    """
    import io
    from datetime import date
    from docx.shared import Pt

    doc = Document(io.BytesIO(modelo_bytes))

    variaveis = {
        "NOME":          funcionario.get("nome", ""),
        "CPF":           funcionario.get("cpf", ""),
        "MATRICULA":     funcionario.get("matricula", funcionario.get("cpf", "")),
        "CARGO":         funcionario.get("cargo", ""),
        "LOTACAO":       funcionario.get("lotacao", ""),
        "DATA_ADMISSAO": funcionario.get("admissao", ""),
        "DATA_HOJE":     date.today().strftime("%d/%m/%Y"),
        "CELULAR":       funcionario.get("celular", ""),
        "EMAIL":         funcionario.get("email", ""),
        "EMPRESA":       EMPRESA,
        "RESP_SST":      RESP_SST,
        "CNPJ":          CNPJ,
        "CTPS":          funcionario.get("ctps", ""),
        "RG":            funcionario.get("rg", ""),
    }

    for table in doc.tables:
        # Primeira passagem: identifica índices das linhas-alvo (cabeçalhos)
        idx_cbo    = None  # linha após "DESCRIÇÃO DAS ATIVIDADES"
        idx_riscos = None  # linha após "RISCOS"
        idx_epis   = None  # linha após "EQUIPAMENTO DE PROTEÇÃO INDIVIDUAL"

        for ri, row in enumerate(table.rows):
            # Pega apenas a primeira célula única para não duplicar
            cells_unicas = []
            seen = set()
            for c in row.cells:
                if id(c._tc) not in seen:
                    seen.add(id(c._tc))
                    cells_unicas.append(c)
            texto_linha = " ".join(c.text.strip() for c in cells_unicas).upper()

            # Cabeçalho: "DESCRIÇÃO DAS ATIVIDADES"
            if ("DESCRI" in texto_linha and "ATIVIDADE" in texto_linha
                    and len(texto_linha) < 50 and idx_cbo is None):
                idx_cbo = ri + 1

            # Cabeçalho: "RISCOS" ou "RISCO A QUE ESTÁ EXPOSTO"
            if ("RISCO" in texto_linha and idx_riscos is None
                    and len(texto_linha) < 60):
                idx_riscos = ri + 1

            # Cabeçalho: "EQUIPAMENTO DE PROTEÇÃO INDIVIDUAL (EPI)"
            if ("EQUIPAMENTO" in texto_linha and "EPI" in texto_linha
                    and "UNIFORME" in texto_linha and idx_epis is None
                    and len(texto_linha) < 80):
                idx_epis = ri + 1

        def _preencher_celula(table, ri, conteudo):
            """Preenche a primeira célula da linha ri com o conteúdo em Arial 8."""
            if ri >= len(table.rows):
                return
            row = table.rows[ri]
            seen = set()
            for c in row.cells:
                if id(c._tc) not in seen:
                    seen.add(id(c._tc))
                    for para in c.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    if c.paragraphs:
                        para = c.paragraphs[0]
                        if para.runs:
                            run = para.runs[0]
                            run.text = conteudo
                        else:
                            run = para.add_run(conteudo)
                        run.font.name = "Arial"
                        run.font.size = Pt(8)
                    break

        if idx_cbo is not None:
            _preencher_celula(table, idx_cbo, descricao_atividades)
        if idx_riscos is not None and riscos_texto:
            _preencher_celula(table, idx_riscos, riscos_texto)
        if idx_epis is not None:
            # EPIs em Arial 8, organizados em 2 colunas
            if idx_epis < len(table.rows):
                row = table.rows[idx_epis]
                seen = set()
                cells_unicas = []
                for c in row.cells:
                    if id(c._tc) not in seen:
                        seen.add(id(c._tc))
                        cells_unicas.append(c)
                cell = cells_unicas[0]
                # Limpa célula
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                # Divide EPIs em 2 colunas
                linhas_epi = [l for l in epis_texto.split("\n") if l.strip()]
                metade = (len(linhas_epi) + 1) // 2
                col1 = linhas_epi[:metade]
                col2 = linhas_epi[metade:]
                max_linhas = max(len(col1), len(col2))
                # Escreve linha a linha intercalando col1 e col2 com tab
                for i in range(max_linhas):
                    txt_c1 = col1[i] if i < len(col1) else ""
                    txt_c2 = col2[i] if i < len(col2) else ""
                    linha_txt = f"{txt_c1:<45}{txt_c2}"
                    if i == 0 and cell.paragraphs and cell.paragraphs[0].runs:
                        run = cell.paragraphs[0].runs[0]
                        run.text = linha_txt
                        run.font.name = "Arial"
                        run.font.size = Pt(8)
                    else:
                        para = cell.add_paragraph()
                        run = para.add_run(linha_txt)
                        run.font.name = "Arial"
                        run.font.size = Pt(8)

        # Segunda passagem: substitui variáveis {{...}} em TODAS as células
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _processar_paragrafo(para, variaveis)

        break  # Processa só a tabela principal

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def criar_os_base_docx() -> bytes:
    """Gera o template base de Ordem de Serviço em DOCX."""
    from docx import Document as _Doc
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    import io as _io

    doc = _Doc()
    # Margens
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    def add_bold(para, text, size=11):
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(size)

    def add_normal(para, text, size=10):
        run = para.add_run(text)
        run.font.size = Pt(size)

    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bold(titulo, "ORDEM DE SERVIÇO - SST", 14)

    doc.add_paragraph()

    # Dados do funcionário
    campos = [
        ("Empresa:", "{{EMPRESA}}"),
        ("CNPJ:", "{{CNPJ}}"),
        ("Nome:", "{{NOME}}"),
        ("CPF:", "{{CPF}}"),
        ("Matrícula:", "{{MATRICULA}}"),
        ("Cargo/Função:", "{{CARGO}}"),
        ("Lotação/Setor:", "{{LOTACAO}}"),
        ("Data de Admissão:", "{{DATA_ADMISSAO}}"),
        ("Celular:", "{{CELULAR}}"),
        ("Data de Emissão:", "{{DATA_HOJE}}"),
    ]
    for label, val in campos:
        p = doc.add_paragraph()
        add_bold(p, f"{label} ")
        add_normal(p, val)

    doc.add_paragraph()

    # Tabela principal com 3 colunas
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    headers = [
        "DESCRIÇÃO DAS ATIVIDADES",
        "RISCOS A QUE ESTÁ EXPOSTO",
        "EQUIPAMENTO DE PROTEÇÃO INDIVIDUAL (EPI) E UNIFORME",
    ]
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)

    placeholders = ["{{ATIVIDADES}}", "{{RISCOS}}", "{{EPIS}}"]
    for i, ph in enumerate(placeholders):
        cell = table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.add_run(ph).font.size = Pt(9)

    doc.add_paragraph()

    # Assinaturas
    p_ass = doc.add_paragraph()
    add_bold(p_ass, "Responsável SST: ")
    add_normal(p_ass, "{{RESP_SST}}")

    p_func = doc.add_paragraph()
    add_bold(p_func, "Assinatura do Funcionário: ")
    add_normal(p_func, "_" * 40)

    p_data = doc.add_paragraph()
    add_bold(p_data, "Data: ")
    add_normal(p_data, "{{DATA_HOJE}}")

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extrair_texto_docx(conteudo_bytes: bytes) -> str:
    """Extrai todo o texto de um .docx em formato editável."""
    import io
    doc = Document(io.BytesIO(conteudo_bytes))
    linhas = []
    # Parágrafos do corpo
    for para in doc.paragraphs:
        linhas.append(para.text)
    # Tabelas
    for tabela in doc.tables:
        linhas.append("")
        for linha in tabela.rows:
            celulas = "\t".join(cel.text.strip() for cel in linha.cells)
            linhas.append(celulas)
    return "\n".join(linhas)


def texto_para_docx(texto: str) -> bytes:
    """Cria um .docx simples a partir de texto (para salvar edições)."""
    import io
    doc = Document()
    for linha in texto.split("\n"):
        # Linha com tabs vira parágrafos separados por espaço (simplificado)
        doc.add_paragraph(linha.replace("\t", "   "))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def pasta_lote() -> str:
    """Cria e retorna pasta para o lote do dia."""
    from datetime import datetime
    nome = datetime.now().strftime("lote_%Y%m%d_%H%M%S")
    pasta = os.path.join(OUTPUT_DIR, nome)
    os.makedirs(pasta, exist_ok=True)
    return pasta


def juntar_pdfs(caminhos_pdf: list[str], pasta_saida: str, nome_funcionario: str) -> str | None:
    """
    Junta múltiplos PDFs em um único arquivo.
    Retorna o caminho do PDF final ou None em caso de erro.
    """
    try:
        import subprocess, re

        nome_seguro = re.sub(r"[^\w\s-]", "", nome_funcionario)
        nome_seguro = re.sub(r"\s+", "_", nome_seguro.strip())
        pdf_final   = os.path.join(pasta_saida, f"Kit_SST__{nome_seguro}.pdf")

        # Tenta usar pdfunite (ghostscript/poppler) — disponível no Docker
        pdfunite = shutil.which("pdfunite")
        if pdfunite:
            cmd = [pdfunite] + caminhos_pdf + [pdf_final]
            resultado = subprocess.run(cmd, capture_output=True, timeout=60)
            if resultado.returncode == 0 and os.path.exists(pdf_final):
                return pdf_final

        # Tenta ghostscript
        gs = shutil.which("gs")
        if gs:
            cmd = [gs, "-dBATCH", "-dNOPAUSE", "-q", "-sDEVICE=pdfwrite",
                   f"-sOutputFile={pdf_final}"] + caminhos_pdf
            resultado = subprocess.run(cmd, capture_output=True, timeout=60)
            if resultado.returncode == 0 and os.path.exists(pdf_final):
                return pdf_final

        # Fallback: usa pypdf
        try:
            from pypdf import PdfWriter
            writer = PdfWriter()
            for caminho in caminhos_pdf:
                writer.append(caminho)
            with open(pdf_final, "wb") as f:
                writer.write(f)
            if os.path.exists(pdf_final):
                return pdf_final
        except ImportError:
            pass

        # Último fallback: retorna o primeiro PDF se não conseguir juntar
        if caminhos_pdf:
            shutil.copy(caminhos_pdf[0], pdf_final)
            return pdf_final

        return None

    except Exception as e:
        print(f"❌ Erro ao juntar PDFs: {e}")
        if caminhos_pdf and os.path.exists(caminhos_pdf[0]):
            try:
                shutil.copy(caminhos_pdf[0], pdf_final)
                return pdf_final
            except Exception:
                return caminhos_pdf[0]
        return None
